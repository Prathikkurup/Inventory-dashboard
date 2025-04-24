from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from flask_cors import CORS
app = Flask(__name__)
CORS(app)




def create_purchase_order_docx(data):
    doc = Document()
    title = doc.add_heading('PURCHASE ORDER', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    po_number = doc.add_heading(f"PO #{data['po_number']}", level=2)
    po_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('BILL TO:', level=3)
    doc.add_paragraph(f"{data['bill_to']['name']}\n{data['bill_to']['address']}\n{data['bill_to']['contact']}")
    
    doc.add_heading('SHIP TO:', level=3)
    doc.add_paragraph(f"{data['ship_to']['name']}\n{data['ship_to']['address']}\n{data['ship_to']['contact']}")
    
    doc.add_paragraph()
    doc.add_heading('Shipping Method and Shipping Terms', level=3)
    doc.add_paragraph(f"Shipping Method: {data['shipping_info']['method']}")
    doc.add_paragraph(f"Ship Via: {data['shipping_info']['via']}")
    doc.add_paragraph(f"Payment: {data['shipping_info']['payment']}")
    doc.add_paragraph(f"Delivery Date: {data['shipping_info']['delivery']}")
    
    doc.add_paragraph()
    doc.add_heading('Order Items', level=3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ITEM NO.'
    hdr_cells[1].text = 'DESCRIPTION'
    hdr_cells[2].text = 'QTY'
    
    for item in data['order_items']:
        row_cells = table.add_row().cells
        row_cells[0].text = item['item_no']
        row_cells[1].text = item['description']
        row_cells[2].text = str(item['qty'])
    
    doc.add_paragraph()
    doc.add_paragraph('Authorized Signature:', style='Heading 3')
    doc.add_paragraph('_')

    file_name = f"Purchase_Order_{data['order_date']}.docx"
    doc.save(file_name)
    return file_name

@app.route('/generate-doc', methods=['POST'])
def generate_doc():
    data = request.json
    file_name = create_purchase_order_docx(data)
    
    return send_file(file_name, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
